import os
from flask import Flask, request, redirect, url_for, render_template, send_from_directory
from werkzeug.utils import secure_filename
import pdfplumber
import pytesseract
from PIL import Image
from googletrans import Translator
from docx import Document

UPLOAD_FOLDER = 'uploads'
DOWNLOAD_FOLDER = 'downloads'
ALLOWED_EXTENSIONS = {'pdf'}

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['DOWNLOAD_FOLDER'] = DOWNLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            try:
                page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
                if page_text and page_text.strip():
                    text += page_text + "\n"
                else:
                    image = page.to_image(resolution=300).original
                    text += pytesseract.image_to_string(image, lang='hin+san+eng') + "\n"
            except Exception as e:
                print(f"Could not process page {page.page_number}: {e}")
                try:
                    image = page.to_image(resolution=300).original
                    text += pytesseract.image_to_string(image, lang='hin+san+eng') + "\n"
                except Exception as ocr_e:
                    print(f"OCR also failed for page {page.page_number}: {ocr_e}")
    return text

def translate_text(text):
    paragraphs = [p.strip() for p in text.splitlines() if p.strip()]
    translator = Translator()
    translated_paragraphs = []
    for para in paragraphs:
        try:
            translated = translator.translate(para, dest='en')
            translated_paragraphs.append({
                'original': para,
                'translated': translated.text
            })
        except Exception as e:
            print(f"Could not translate paragraph: {e}")
            translated_paragraphs.append({
                'original': para,
                'translated': '[Translation Error]'
            })
    return translated_paragraphs

def create_docx(content, original_filename):
    doc = Document()
    doc.add_heading(f'Translation of {original_filename}', 0)

    for item in content:
        p_original = doc.add_paragraph(item['original'])
        p_original.style.font.italic = True
        p_original.style.font.color.rgb = (80, 80, 80)
        doc.add_paragraph(item['translated'])
        doc.add_paragraph() # Add a space between translations

    docx_filename = f"translated_{os.path.splitext(original_filename)[0]}.docx"
    filepath = os.path.join(app.config['DOWNLOAD_FOLDER'], docx_filename)
    doc.save(filepath)
    return docx_filename

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            extracted_text = extract_text_from_pdf(filepath)
            translated_content = translate_text(extracted_text)
            docx_filename = create_docx(translated_content, filename)

            return render_template('result.html',
                                   content=translated_content,
                                   docx_filename=docx_filename)

    return render_template('index.html')

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(app.config['DOWNLOAD_FOLDER'], filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)